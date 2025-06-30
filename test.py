import os
import re
import threading
import tkinter as tk
from tkinter import simpledialog, messagebox
from tkinter import filedialog
from docx import Document
import pyautogui
import pytesseract
from PIL import Image, ImageTk
import pystray
import keyboard
import sys
import time

# ========== 配置区 ==========
def find_tesseract():
    base = getattr(sys, '_MEIPASS', os.path.abspath('.'))
    local_path = os.path.join(base, 'Tesseract-OCR', 'tesseract.exe')
    if os.path.exists(local_path):
        return local_path
    candidates = [
        r'C:/Program Files/Tesseract-OCR/tesseract.exe',
        r'C:/Program Files (x86)/Tesseract-OCR/tesseract.exe',
        r'D:/Tesseract-OCR/tesseract.exe'
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    root = tk.Tk()
    root.withdraw()
    path = filedialog.askopenfilename(title='请选择tesseract.exe', filetypes=[('exe', '*.exe')])
    return path

tesseract_path = find_tesseract()
if tesseract_path:
    pytesseract.pytesseract.tesseract_cmd = tesseract_path
else:
    raise RuntimeError('未找到Tesseract-OCR，请安装或手动指定路径')

# ========== 记忆上次目录 ==========
last_folder = {'path': None}

def get_folder():
    if last_folder['path'] and os.path.exists(last_folder['path']):
        return last_folder['path']
    folder = filedialog.askdirectory(title='选择要搜索的文件夹')
    if folder:
        last_folder['path'] = folder
    return folder

def set_folder(folder):
    if folder:
        last_folder['path'] = folder

# ========== 屏幕区域选择 ==========
def select_screen_area(master):
    sel = {'x1': 0, 'y1': 0, 'x2': 0, 'y2': 0, 'done': False}
    top = tk.Toplevel(master)
    top.attributes('-alpha', 0.3)
    top.attributes('-topmost', True)
    top.overrideredirect(True)
    top.geometry(f"{top.winfo_screenwidth()}x{top.winfo_screenheight()}+0+0")
    canvas = tk.Canvas(top, cursor='cross', bg='gray')
    canvas.pack(fill='both', expand=True)
    rect = [None]

    def on_press(event):
        sel['x1'], sel['y1'] = event.x, event.y
        if rect[0]:
            canvas.delete(rect[0])
        rect[0] = canvas.create_rectangle(sel['x1'], sel['y1'], sel['x1'], sel['y1'], outline='red', width=2)

    def on_drag(event):
        sel['x2'], sel['y2'] = event.x, event.y
        canvas.coords(rect[0], sel['x1'], sel['y1'], sel['x2'], sel['y2'])

    def on_release(event):
        sel['x2'], sel['y2'] = event.x, event.y
        sel['done'] = True
        top.destroy()

    canvas.bind('<Button-1>', on_press)
    canvas.bind('<B1-Motion>', on_drag)
    canvas.bind('<ButtonRelease-1>', on_release)
    master.wait_window(top)
    x1, y1, x2, y2 = sel['x1'], sel['y1'], sel['x2'], sel['y2']
    left, top_ = min(x1, x2), min(y1, y2)
    width, height = abs(x2 - x1), abs(y2 - y1)
    return left, top_, width, height

# ========== 悬浮窗主类 ==========
class FloatingWindow:
    def __init__(self, root):
        self.root = root
        self.root.overrideredirect(True)
        self.root.attributes('-topmost', True)
        self.root.attributes('-alpha', 0.85)
        self.root.geometry(f'300x60+{self.root.winfo_screenwidth()-320}+20')
        self.root.withdraw()  # 启动时隐藏
        self.drag_data = {'x': 0, 'y': 0}
        self.create_widgets()
        self.bind_events()

    def create_widgets(self):
        self.frame = tk.Frame(self.root, bg='#F0F0F0', bd=2, relief='ridge')
        self.frame.pack(fill='both', expand=True)
        self.search_btn = tk.Button(self.frame, text='🔍手动搜索', command=self.manual_search)
        self.search_btn.pack(side='left', padx=8, pady=10)
        self.ocr_btn = tk.Button(self.frame, text='🖼️屏幕取词', command=self.ocr_search)
        self.ocr_btn.pack(side='left', padx=8, pady=10)
        self.exit_btn = tk.Button(self.frame, text='❌', command=self.hide)
        self.exit_btn.pack(side='right', padx=8, pady=10)

    def bind_events(self):
        self.frame.bind('<Button-1>', self.start_move)
        self.frame.bind('<B1-Motion>', self.on_motion)
        self.root.bind('<Escape>', lambda e: self.hide())

    def start_move(self, event):
        self.drag_data['x'] = event.x
        self.drag_data['y'] = event.y

    def on_motion(self, event):
        x = self.root.winfo_pointerx() - self.drag_data['x']
        y = self.root.winfo_pointery() - self.drag_data['y']
        self.root.geometry(f'+{x}+{y}')

    def show(self):
        self.root.deiconify()
        self.root.lift()

    def hide(self):
        self.root.withdraw()

    def manual_search(self):
        keyword = simpledialog.askstring('输入关键词', '请输入要搜索的内容（支持正则表达式）：', parent=self.root)
        if not keyword:
            return
        folder = get_folder()
        if not folder:
            return
        set_folder(folder)
        result = search_in_docx(keyword, folder)
        if result:
            msg = '\n'.join(result)
            messagebox.showinfo('搜索结果', msg[:2000] + ('...\n结果过多已截断' if len(msg)>2000 else ''))
        else:
            messagebox.showinfo('搜索结果', '未找到匹配内容')

    def ocr_search(self):
        self.hide()
        time.sleep(0.2)
        messagebox.showinfo('提示', '请用鼠标拖拽框选需要识别的屏幕区域')
        left, top, width, height = select_screen_area(self.root)
        print(f"框选区域: left={left}, top={top}, width={width}, height={height}")  # 调试用
        if width == 0 or height == 0:
            messagebox.showinfo('OCR结果', '未框选区域')
            self.show()
            return
        time.sleep(0.3)
        img = pyautogui.screenshot(region=(left, top, width, height))
        img.save('debug_ocr.png')
        text = pytesseract.image_to_string(img, lang='chi_sim')
        print(f"OCR识别结果: {text}")
        if not text.strip():
            messagebox.showinfo('OCR结果', '未识别到文字')
            self.show()
            return
        # 去除所有空格
        clean_text = text.replace(' ', '').replace('\u3000', '').replace('\t', '').replace('\n', '')
        # 复制到剪贴板
        self.root.clipboard_clear()
        self.root.clipboard_append(clean_text)
        self.root.update()  # 保证剪贴板内容可用
        folder = get_folder()
        if not folder:
            self.show()
            return
        set_folder(folder)
        result = search_in_docx(clean_text, folder)
        if result:
            msg = '\n'.join(result)
            messagebox.showinfo('搜索结果', msg[:2000] + ('...\n结果过多已截断' if len(msg)>2000 else ''))
        else:
            messagebox.showinfo('搜索结果', '未找到匹配内容')
        self.show()

# ========== Word文档搜索 ==========
def search_in_docx(keyword, folder):
    result = []
    try:
        pattern = re.compile(keyword)
    except Exception as e:
        messagebox.showerror('正则错误', f'正则表达式有误: {e}')
        return result
    for root, dirs, files in os.walk(folder):
        for file in files:
            if file.lower().endswith('.docx'):
                path = os.path.join(root, file)
                try:
                    doc = Document(path)
                    for para in doc.paragraphs:
                        if pattern.search(para.text):
                            result.append(f'{file}: {para.text.strip()}')
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if pattern.search(cell.text):
                                    result.append(f'{file}: {cell.text.strip()}')
                except Exception as e:
                    result.append(f'{file}: 读取失败({e})')
    return result

# ========== 托盘图标 ==========
def create_tray(window: FloatingWindow):
    def on_show(icon, item):
        window.show()
    def on_exit(icon, item):
        icon.stop()
        window.root.quit()
        os._exit(0)  # 强制退出所有线程和主循环
    image = Image.new('RGB', (64, 64), color=(70, 130, 180))
    icon = pystray.Icon('WordSearch', image, 'Word文档搜索', menu=pystray.Menu(
        pystray.MenuItem('显示窗口', on_show),
        pystray.MenuItem('退出', on_exit)
    ))
    threading.Thread(target=icon.run, daemon=True).start()

# ========== 全局热键 ==========
def hotkey_listener(window: FloatingWindow):
    def toggle():
        if window.root.state() == 'withdrawn':
            window.show()
        else:
            window.hide()
    keyboard.add_hotkey('ctrl+alt+f', toggle)
    keyboard.wait()  # 阻塞线程

# ========== 主程序入口 ==========
def main():
    root = tk.Tk()
    window = FloatingWindow(root)
    create_tray(window)
    threading.Thread(target=hotkey_listener, args=(window,), daemon=True).start()
    root.mainloop()

if __name__ == '__main__':
    main()
